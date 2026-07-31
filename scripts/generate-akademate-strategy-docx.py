#!/usr/bin/env python3
"""Build the Akademate public platform strategy document from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/plans/2026-07-31-akademate-public-platform-expansion.md"
OUTPUT = ROOT / "docs/strategy/Akademate_Strategy_Reservations_Verticals_Domains_Payments.docx"
HERO = ROOT / "apps/web/public/images/marketing/akademate-multisite-network.jpg"

NAVY = RGBColor(7, 22, 51)
BLUE = RGBColor(37, 99, 235)
SLATE = RGBColor(71, 85, 105)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [[part.strip() for part in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for index, row in enumerate(rows) if index != 1]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Shading Accent 1"
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value.replace("**", "")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = NAVY
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                set_cell_shading(cell, "EAF1FF")
    document.add_paragraph()


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = SLATE
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in (("Title", 34), ("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 12)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = NAVY
        style.font.bold = True


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("AKADEMATE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Strategy: Reservations, Verticals, Domains and Payments")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Public SaaS platform direction · 31 July 2026")
    run.font.size = Pt(13)
    run.font.color.rgb = SLATE

    if HERO.exists():
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(HERO), width=Inches(6.6))

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Derived from the ingested strategy conversation and implemented public-web plan. "
        "Scope is Akademate SaaS only; CEP production, data and deployment are excluded."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE
    document.add_section(WD_SECTION.NEW_PAGE)


def add_body(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line or line.startswith("# Akademate Public Platform Expansion"):
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            table_lines = [line]
            index += 1
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            document.add_heading(heading.group(2).replace("**", ""), level=level)
            index += 1
            continue

        checklist = re.match(r"^- \[([ xX])\]\s+(.+)$", line)
        bullet = re.match(r"^-\s+(.+)$", line)
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if checklist:
            marker = "☑" if checklist.group(1).lower() == "x" else "☐"
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{marker} {checklist.group(2)}")
        elif bullet:
            document.add_paragraph(bullet.group(1).replace("**", ""), style="List Bullet")
        elif numbered:
            document.add_paragraph(numbered.group(1).replace("**", ""), style="List Number")
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(line.replace("**", "").replace("`", ""))
        index += 1


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Akademate public SaaS strategy · SOLARIA AGENCY OÜ · Informational working document"
        run = footer.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = SLATE


def main() -> None:
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    add_cover(document)
    add_body(document, SOURCE.read_text(encoding="utf-8"))
    add_footer(document)
    document.core_properties.title = "Akademate Strategy: Reservations, Verticals, Domains and Payments"
    document.core_properties.subject = "Public SaaS operating model and implementation plan"
    document.core_properties.author = "Akademate / SOLARIA AGENCY OÜ"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
